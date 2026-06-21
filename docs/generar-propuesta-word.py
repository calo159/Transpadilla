"""
Genera PROPUESTA-ALCALDIA.docx a partir de PROPUESTA-ALCALDIA.md.

El Markdown es la ÚNICA fuente de verdad: edita el .md y vuelve a ejecutar
  python generar-propuesta-word.py
para regenerar el Word (queda junto a este script, en docs/).

Soporta: títulos (#..####), párrafos, **negrita**, *cursiva*, `código`,
enlaces [texto](url), listas con viñeta y numeradas, citas (>), reglas (---)
y tablas estilo Markdown (| a | b |).
"""
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

AZUL = RGBColor(0x0D, 0x24, 0x61)   # azul marino TransPadilla
GRIS = RGBColor(0x55, 0x60, 0x70)

AQUI = os.path.dirname(os.path.abspath(__file__))
ENTRADA = os.path.join(AQUI, "PROPUESTA-ALCALDIA.md")
SALIDA = os.path.join(AQUI, "PROPUESTA-ALCALDIA.docx")

INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))")


def add_inline(p, texto):
    """Agrega texto a un párrafo interpretando negrita/cursiva/código/enlaces."""
    for parte in INLINE.split(texto):
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**"):
            r = p.add_run(parte[2:-2]); r.bold = True
        elif parte.startswith("`") and parte.endswith("`"):
            r = p.add_run(parte[1:-1]); r.font.name = "Consolas"
        elif parte.startswith("*") and parte.endswith("*"):
            r = p.add_run(parte[1:-1]); r.italic = True
        elif parte.startswith("[") and "](" in parte:
            txt = parte[1:parte.index("](")]
            p.add_run(txt)
        else:
            p.add_run(parte)


def fila_tabla(linea):
    return [c.strip() for c in linea.strip().strip("|").split("|")]


def es_separador(linea):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", linea)) and "-" in linea


def main():
    with open(ENTRADA, encoding="utf-8") as f:
        lineas = f.read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i, n = 0, len(lineas)
    while i < n:
        s = lineas[i].strip()

        # Tabla
        if s.startswith("|") and i + 1 < n and es_separador(lineas[i + 1]):
            encabezados = fila_tabla(s)
            i += 2
            filas = []
            while i < n and lineas[i].strip().startswith("|"):
                filas.append(fila_tabla(lineas[i]))
                i += 1
            tabla = doc.add_table(rows=1, cols=len(encabezados))
            tabla.style = "Light Grid Accent 1"
            for j, h in enumerate(encabezados):
                par = tabla.rows[0].cells[j].paragraphs[0]
                par.clear()
                add_inline(par, h)
                for run in par.runs:
                    run.bold = True
            for fila in filas:
                celdas = tabla.add_row().cells
                for j, val in enumerate(fila[: len(encabezados)]):
                    par = celdas[j].paragraphs[0]
                    par.clear()
                    add_inline(par, val)
            doc.add_paragraph()
            continue

        if s == "" or s == "---":
            i += 1
            continue

        # Títulos
        if s.startswith("#"):
            nivel = len(s) - len(s.lstrip("#"))
            h = doc.add_heading(level=min(nivel, 4))
            h.clear()
            add_inline(h, s[nivel:].strip())
            for run in h.runs:
                run.font.color.rgb = AZUL
            i += 1
            continue

        # Cita
        if s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_inline(p, s.lstrip("> ").strip())
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = GRIS
            i += 1
            continue

        # Listas
        if re.match(r"^[-*]\s+", s):
            add_inline(doc.add_paragraph(style="List Bullet"), re.sub(r"^[-*]\s+", "", s))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", s):
            add_inline(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s+", "", s))
            i += 1
            continue

        # Párrafo normal
        add_inline(doc.add_paragraph(), s)
        i += 1

    doc.add_paragraph()
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pie.add_run("TransPadilla  ·  Moviendo la Ciudad  ·  Riohacha, La Guajira")
    r.bold = True
    r.font.color.rgb = AZUL

    doc.save(SALIDA)
    print(f"OK Word generado: {SALIDA}")


if __name__ == "__main__":
    main()
