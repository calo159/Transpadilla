"""
Genera PROPUESTA-ALCALDIA.docx — documento Word profesional para la Alcaldía.
Uso: python generar-propuesta-word.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores corporativos ──────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1B, 0x4F, 0x8A)   # encabezados y tablas
AZUL_CLARO   = RGBColor(0x4B, 0xA9, 0xD8)   # portada / acento
GRIS_FILA    = RGBColor(0xE8, 0xF0, 0xF7)   # filas alternas de tabla
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
NEGRO        = RGBColor(0x1A, 0x1A, 0x2E)


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color: RGBColor):
    """Relleno de fondo de celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    """Borde fino en todas las caras de la celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def bold_run(para, text, size=None, color=None, italic=False):
    run = para.add_run(text)
    run.bold = True
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def normal_run(para, text, size=None, color=None, italic=False):
    run = para.add_run(text)
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_heading1(doc, text):
    h = doc.add_heading(level=1)
    h.clear()
    run = h.add_run(text)
    run.font.color.rgb = AZUL_OSCURO
    run.font.size = Pt(14)
    run.bold = True
    run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_heading2(doc, text):
    h = doc.add_heading(level=2)
    h.clear()
    run = h.add_run(text)
    run.font.color.rgb = AZUL_OSCURO
    run.font.size = Pt(12)
    run.bold = True
    run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_body(doc, text, bold_parts=None):
    """Párrafo de cuerpo. bold_parts: lista de strings que van en negrita."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_parts:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.name = "Calibri"
                r.font.size = Pt(11)
            r2 = p.add_run(bp)
            r2.bold = True
            r2.font.name = "Calibri"
            r2.font.size = Pt(11)
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.name = "Calibri"
            r.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
        rest = text[len(bold_prefix):]
        r2 = p.add_run(rest)
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
    return p


def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
        rest = text[len(bold_prefix):]
        r2 = p.add_run(rest)
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def make_table(doc, headers, rows, col_widths=None):
    """Tabla con encabezado azul oscuro y filas alternas."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"

    # Encabezado
    hdr = t.rows[0]
    for i, htext in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, AZUL_OSCURO)
        set_cell_borders(cell, "1B4F8A")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(htext)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.name = "Calibri"
        run.font.size = Pt(10)

    # Filas de datos
    for r_idx, row_data in enumerate(rows):
        row = t.rows[r_idx + 1]
        bg = GRIS_FILA if r_idx % 2 == 0 else BLANCO
        for c_idx, ctext in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Negritas inline: texto entre ** **
            parts = ctext.split("**")
            for pi, part in enumerate(parts):
                r = p.add_run(part)
                r.bold = (pi % 2 == 1)
                r.font.name = "Calibri"
                r.font.size = Pt(10)

    # Anchos de columna
    if col_widths:
        for c_idx, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[c_idx].width = Cm(w)

    doc.add_paragraph()  # espacio tras la tabla
    return t


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(
        "TransPadilla — Moviendo la Ciudad  ·  Riohacha, La Guajira  ·  Documento Confidencial"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def add_page_break(doc):
    doc.add_page_break()


# ── Documento ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.5)

    # Estilo Normal base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_footer(doc)

    # ── PORTADA ──────────────────────────────────────────────────────────────
    # Espacio superior
    for _ in range(6):
        doc.add_paragraph()

    # Logo / Nombre
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("TransPadilla")
    r.font.name = "Calibri"
    r.font.size = Pt(40)
    r.bold = True
    r.font.color.rgb = AZUL_OSCURO

    # Subtítulo
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("Sistema de Rastreo de Transporte Público en Tiempo Real")
    r2.font.name = "Calibri"
    r2.font.size = Pt(16)
    r2.font.color.rgb = AZUL_CLARO

    doc.add_paragraph()

    # Etiqueta documento
    label_p = doc.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = label_p.add_run("Propuesta de Implementación para la Alcaldía de Riohacha")
    r3.font.name = "Calibri"
    r3.font.size = Pt(13)
    r3.bold = True
    r3.font.color.rgb = NEGRO

    # Fecha
    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = date_p.add_run("Riohacha, La Guajira  ·  Junio 2026")
    r4.font.name = "Calibri"
    r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Línea decorativa
    for _ in range(4):
        doc.add_paragraph()
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = line_p.add_run("─" * 55)
    r5.font.color.rgb = AZUL_CLARO
    r5.font.size = Pt(12)

    add_page_break(doc)

    # ── SECCIÓN 1 — RESUMEN EJECUTIVO ────────────────────────────────────────
    add_heading1(doc, "1. Resumen Ejecutivo")
    add_body(doc,
        "TransPadilla permite a la ciudadanía ver en vivo dónde están los buses, cuánto "
        "tardan en llegar y el estado del servicio; a los conductores reportar su recorrido, "
        "ocupación y novedades; y a la Alcaldía supervisar la flota y el estado del tráfico "
        "en tiempo real.",
        bold_parts=["en vivo", "Alcaldía supervisar la flota"]
    )
    add_body(doc,
        "El sistema ya está construido, probado y funcionando en una versión de demostración "
        "accesible desde cualquier dispositivo con conexión a internet. Para operar de forma "
        "continua y confiable las 24 horas del día, los 7 días de la semana, se requiere una "
        "inversión modesta en infraestructura y operación, detallada en este documento.",
        bold_parts=["ya está construido, probado y funcionando", "24 horas del día, los 7 días de la semana"]
    )
    add_body(doc,
        "Esta inversión garantiza alta disponibilidad, detección temprana de fallos, respaldo "
        "de datos y mantenimiento para corregir rápidamente cualquier incidencia.",
        bold_parts=["alta disponibilidad", "detección temprana de fallos", "respaldo de datos", "mantenimiento"]
    )

    # ── SECCIÓN 2 — MEJORAS YA REALIZADAS ────────────────────────────────────
    add_heading1(doc, "2. Mejoras Ya Realizadas para Producción (Sin Costo Adicional)")
    add_body(doc,
        "Se implementó, sin costo adicional, el endurecimiento técnico necesario para que el "
        "sistema sea apto para una entidad pública:"
    )

    bullets = [
        ("Seguridad: ", "secretos obligatorios en producción, cabeceras de seguridad, límite de "
            "intentos de inicio de sesión (anti fuerza bruta), validación de todos los datos de "
            "entrada y restricción de orígenes (CORS)."),
        ("Robustez: ", "manejo global de errores, verificación automática de la base de datos, "
            "apagado ordenado para reinicios 24/7 y pantalla de recuperación ante fallos (sin "
            "\"pantallas blancas\")."),
        ("Datos de producción limpios: ", "arranque configurable sin datos de prueba, con solo "
            "el administrador real de la Alcaldía."),
        ("Flexibilidad de proveedores: ", "el proveedor de mapas y cálculo de rutas se puede "
            "cambiar sin reprogramar el sistema (solo configuración)."),
        ("GPS del conductor: ", "la pantalla del teléfono se mantiene encendida durante el "
            "recorrido (Wake Lock) y el GPS se reactiva automáticamente al volver a la "
            "aplicación, garantizando transmisión continua."),
        ("Operación autocontenida: ", "archivos Docker y guía de despliegue completa para "
            "hospedar todo el sistema en un único servidor económico, con respaldos y monitoreo."),
    ]
    for prefix, rest in bullets:
        add_bullet(doc, prefix + rest, bold_prefix=prefix)

    # ── SECCIÓN 3 — FINANCIAMIENTO ───────────────────────────────────────────
    add_heading1(doc, "3. Inversión Requerida para Operación 24/7")

    # 3.1
    add_heading2(doc, "3.1 Infraestructura (Hospedaje) — Obligatorio")
    add_body(doc,
        "La versión de demostración usa un plan gratuito que se apaga por inactividad. "
        "Para operación 24/7 se requiere hospedaje pagado. Dos alternativas:"
    )
    make_table(doc,
        ["Alternativa", "Descripción", "Costo mensual aprox."],
        [
            ["**VPS único (recomendado)**",
             "Un servidor (Hetzner / DigitalOcean) con Docker: base de datos + Django + web",
             "US$6–15\n(~$24.000–60.000 COP)"],
            ["Render (servicios gestionados)",
             "Web + Django + PostgreSQL con backups automáticos",
             "US$20–50\n(~$80.000–200.000 COP)"],
        ],
        col_widths=[5, 9, 4.5]
    )

    # 3.2
    add_heading2(doc, "3.2 Mapas y Cálculo de Rutas — Recomendado")
    add_body(doc,
        "Los servicios públicos actuales (OpenStreetMap / OSRM de demostración) no tienen "
        "garantía de disponibilidad para uso institucional intensivo."
    )
    make_table(doc,
        ["Ítem", "Opción", "Costo mensual aprox."],
        [
            ["Mapa (tiles / imágenes)", "MapTiler o Mapbox (capa gratuita disponible)", "US$0–50\n(~$0–200.000 COP)"],
            ["Cálculo de rutas", "OSRM propio en el mismo servidor VPS", "**US$0** (incluido)"],
        ],
        col_widths=[5, 8, 4.5]
    )

    # 3.3
    add_heading2(doc, "3.3 GPS de los Buses — Decisión Principal")
    add_body(doc,
        "Esta es la decisión más importante: cómo reporta su posición cada bus. "
        "El reto es la transmisión continua, especialmente con la pantalla apagada.",
        bold_parts=["cómo reporta su posición cada bus"]
    )
    add_note(doc,
        "Nota técnica: una aplicación web no puede transmitir con la pantalla totalmente "
        "apagada (limitación de los navegadores). Ya se implementó sin costo que la "
        "pantalla se mantenga encendida durante el recorrido (Wake Lock), cubriendo el "
        "caso más común. Para transmisión 100% en segundo plano se requiere app nativa "
        "o rastreador dedicado."
    )
    make_table(doc,
        ["Opción", "Confiabilidad", "Inversión inicial", "Costo mensual / bus", "¿Depende del conductor?"],
        [
            ["A) Web actual + Wake Lock (ya hecho)",
             "Media\n(pantalla encendida)",
             "Smartphone si no cuentan con uno",
             "Plan de datos\nUS$5–10",
             "Sí\n(debe dejar la app abierta)"],
            ["B) App nativa Android\n(Capacitor)",
             "Alta\n(GPS en segundo plano)",
             "~US$25 cuenta Google Play (una vez)\n+ plugin opcional ~US$300",
             "Plan de datos\nUS$5–10",
             "Sí\n(debe llevar el teléfono)"],
            ["**C) Rastreador GPS dedicado (recomendado)**",
             "**Máxima**\n(transmite solo)",
             "Equipo US$30–80 por bus\n(pago único)",
             "SIM datos\nUS$3–8",
             "**No**\n(opera de forma autónoma)"],
        ],
        col_widths=[4.5, 3, 4.5, 3.5, 3]
    )
    add_note(doc,
        "Ejemplo con flota de 20 buses y rastreadores dedicados: inversión inicial "
        "~US$600–1.600 (~$2,4M–6,4M COP) + mensual ~US$60–160 (~$240.000–640.000 COP). "
        "Recomendación: para operación institucional 24/7 el rastreador dedicado es la "
        "opción más confiable (no depende del comportamiento del conductor). La app nativa "
        "es un buen punto medio si se desea evitar compra de hardware."
    )

    # 3.4
    add_heading2(doc, "3.4 Confiabilidad y Operación — Recomendado")
    make_table(doc,
        ["Ítem", "Para qué sirve", "Costo aprox."],
        [
            ["Dominio propio (.gov.co / .com)", "Imagen institucional y acceso fácil", "US$10–40 / año"],
            ["Monitoreo de uptime (UptimeRobot)", "Alerta inmediata si el sistema cae", "US$0–10 / mes"],
            ["Rastreo de errores (Sentry)", "Detectar fallos técnicos en vivo", "US$0–26 / mes"],
            ["Respaldos de base de datos", "Proteger la información de la flota", "Incluido en VPS"],
        ],
        col_widths=[5.5, 7.5, 4]
    )

    # 3.5
    add_heading2(doc, "3.5 Mantenimiento y Soporte — Clave para \"Sin Errores\"")
    add_body(doc,
        "Todo sistema 24/7 requiere mantenimiento continuo: parches de seguridad, "
        "actualizaciones, soporte y monitoreo. Es lo que sostiene la confiabilidad "
        "a largo plazo. Se cotiza por horas de desarrollador o como contrato de "
        "soporte mensual, según el alcance que defina la Alcaldía.",
        bold_parts=["mantenimiento continuo", "contrato de soporte mensual"]
    )

    # ── SECCIÓN 4 — PRESUPUESTO ──────────────────────────────────────────────
    add_heading1(doc, "4. Presupuesto Consolidado")

    add_heading2(doc, "Costos únicos (pago inicial)")
    make_table(doc,
        ["Concepto", "Estimado"],
        [
            ["Dominio (primer año)", "~US$15 (~$60.000 COP)"],
            ["Rastreadores GPS (por bus, opcional)", "~US$30–80 c/u"],
            ["Configuración inicial / puesta en marcha", "A acordar (ver soporte)"],
        ],
        col_widths=[12, 6]
    )

    add_heading2(doc, "Costos recurrentes mensuales — escenario económico (VPS)")
    make_table(doc,
        ["Concepto", "Estimado mensual"],
        [
            ["Hospedaje (VPS único)", "US$6–15"],
            ["Mapas con SLA (opcional)", "US$0–50"],
            ["Monitoreo + rastreo de errores", "US$0–30"],
            ["Plan de datos GPS por bus", "US$5–10 × nº de buses"],
            ["**Núcleo del sistema (sin GPS ni soporte)**", "**~US$10–60/mes** (~$40.000–240.000 COP)"],
        ],
        col_widths=[12, 6]
    )
    add_note(doc,
        "El mantenimiento y soporte se acuerda por separado; es lo que garantiza la "
        "operación continua y la corrección rápida de cualquier incidencia."
    )

    # ── SECCIÓN 5 — PLAN DE IMPLEMENTACIÓN ───────────────────────────────────
    add_heading1(doc, "5. Plan de Implementación Sugerido")

    items = [
        ("Fase 1 — Puesta en marcha (1–2 semanas): ",
         "contratar VPS y dominio, desplegar el sistema, activar HTTPS, respaldos y "
         "monitoreo. Cargar las rutas y paradas reales de Riohacha. Crear las cuentas "
         "de administrador y conductores."),
        ("Fase 2 — Piloto (2–4 buses): ",
         "equipar las primeras unidades con GPS (celular o rastreador), capacitar a los "
         "conductores, validar el sistema en campo y ajustar según los resultados."),
        ("Fase 3 — Escalado a la flota completa: ",
         "equipar todos los buses, afinar el sistema con datos reales y comunicar el "
         "servicio a la ciudadanía de Riohacha."),
        ("Continuo — Operación y mantenimiento: ",
         "monitoreo 24/7, respaldos automáticos, soporte y actualizaciones periódicas."),
    ]
    for prefix, rest in items:
        add_numbered(doc, prefix + rest, bold_prefix=prefix)

    # ── SECCIÓN 6 — CONCLUSIÓN ───────────────────────────────────────────────
    add_heading1(doc, "6. Conclusión")
    add_body(doc,
        "El producto ya existe y funciona. La inversión solicitada es principalmente "
        "operativa (hospedaje, GPS de los buses y soporte), no de desarrollo desde cero. "
        "Con un presupuesto modesto, la Alcaldía de Riohacha puede ofrecer a su ciudadanía "
        "un servicio de transporte moderno, transparente y supervisable en tiempo real.",
        bold_parts=["ya existe y funciona", "moderno, transparente y supervisable en tiempo real"]
    )
    add_body(doc,
        "TransPadilla está listo para dar el siguiente paso. Quedamos a disposición "
        "para ampliar información, realizar una demostración en vivo o ajustar cualquier "
        "aspecto de esta propuesta."
    )

    # Firma final
    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = firma.add_run("TransPadilla  ·  Moviendo la Ciudad  ·  Riohacha, La Guajira")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = AZUL_OSCURO
    r.bold = True

    # ── Guardar ───────────────────────────────────────────────────────────────
    out = "PROPUESTA-ALCALDIA.docx"
    doc.save(out)
    print(f"OK Documento generado: {out}")


if __name__ == "__main__":
    build()
